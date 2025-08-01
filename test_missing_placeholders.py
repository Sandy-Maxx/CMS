#!/usr/bin/env python3

import sys
import os
import logging
from docx import Document

# Configure logging to see debug output
logging.basicConfig(
    level=logging.DEBUG,
    format='%(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)

# Add current directory to Python path
sys.path.append('.')

from features.AutodocGen.data_fetcher import DataFetcher
from features.AutodocGen.document_generator import DocumentGenerator
from config import DATABASE_PATH

def test_missing_placeholders():
    """Test placeholder replacement with placeholders that don't exist in database"""
    
    print("Creating a test template with missing placeholders...")
    
    # Create a test template with placeholders that won't be found
    test_template = Document()
    test_template.add_paragraph("Test document with missing placeholders:")
    test_template.add_paragraph("LOA No: [LOA_NO]")  # This should work
    test_template.add_paragraph("Missing Field: [MISSING_FIELD]")  # This should trigger the issue
    test_template.add_paragraph("Tender Opening Date: [TENDER_OPENING_DATE]")  # This should work
    test_template.add_paragraph("Another Missing: [ANOTHER_MISSING]")  # This should trigger the issue
    test_template.save("test_template_missing.docx")
    
    # Initialize the data fetcher and document generator
    try:
        data_fetcher = DataFetcher(DATABASE_PATH)
        generator = DocumentGenerator(data_fetcher)
        
        # Test data
        test_data = {
            'work_id': 1,
            'firm_name': 'Test Firm'
        }
        
        # Template and output paths
        template_path = 'test_template_missing.docx'
        output_path = 'debug_output_missing.docx'
        
        print(f"Using template: {template_path}")
        print(f"Output will be saved to: {output_path}")
        print(f"Test data: {test_data}")
        
        # Generate the document
        generator.generate(template_path, test_data, output_path, is_firm_specific=False)
        
        print("Document generation completed!")
        
        # Read the output to see what happened to the placeholders
        print("\n--- OUTPUT DOCUMENT CONTENT ---")
        try:
            output_doc = Document(output_path)
            for i, paragraph in enumerate(output_doc.paragraphs):
                text = paragraph.text
                if text.strip():
                    print(f"Paragraph {i}: {repr(text)}")
                    
        except Exception as e:
            print(f"Error reading output document: {e}")
        
    except Exception as e:
        print(f"Error during test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_missing_placeholders()
