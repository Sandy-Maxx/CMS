"""
Detailed investigation of placeholder replacement issues.

This test will help us understand why certain placeholders are not being replaced correctly.
"""

import os
import tempfile
import sqlite3
from docx import Document
from features.template_engine.template_processor import TemplateProcessor
from features.template_engine.work_data_provider import WorkDataProvider
from config import DATABASE_PATH
import config
import database.managers.database_utils as db_utils

def investigate_placeholder_replacement():
    print("=== PLACEHOLDER REPLACEMENT INVESTIGATION ===\n")
    
    # Create a temporary database with test data
    with tempfile.NamedTemporaryFile(suffix='.db', delete=False) as temp_db:
        temp_db_path = temp_db.name
    
    try:
        # Set up test database
        conn = sqlite3.connect(temp_db_path)
        cursor = conn.cursor()
        
        # Create works table with test_col
        cursor.execute("""
            CREATE TABLE works (
                id INTEGER PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT,
                test_col TEXT
            )
        """)
        
        # Create firm_documents table
        cursor.execute("""
            CREATE TABLE firm_documents (
                id INTEGER PRIMARY KEY,
                work_id INTEGER,
                firm_name TEXT,
                pg_no TEXT,
                pg_amount REAL,
                bank_name TEXT,
                bank_address TEXT,
                submission_date TEXT,
                pg_submitted INTEGER
            )
        """)
        
        # Insert test data
        cursor.execute(
            "INSERT INTO works (name, description, test_col) VALUES (?, ?, ?)",
            ("Test Work", "Test Description", "Test Column Value")
        )
        
        cursor.execute(
            "INSERT INTO firm_documents (work_id, firm_name, pg_no, pg_amount, bank_name, bank_address, submission_date, pg_submitted) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (1, "Test Firm", "PG123", 50000.0, "Test Bank", "Test Address", "2024-01-01", 1)
        )
        
        conn.commit()
        conn.close()
        
        # Redirect database to test database
        original_database_path = DATABASE_PATH
        config.DATABASE_PATH = temp_db_path
        db_utils.DATABASE_PATH = temp_db_path
        
        try:
            print("1. Testing WorkDataProvider with test database...")
            work_data_provider = WorkDataProvider(1)
            
            # Test work details retrieval
            print(f"Work details: {work_data_provider.work_details}")
            print(f"Firm documents: {work_data_provider.firm_documents}")
            
            # Test placeholder generation
            placeholders = work_data_provider.generate_placeholders()
            print(f"\nGenerated placeholders ({len(placeholders)} total):")
            for key, value in sorted(placeholders.items()):
                print(f"  {key}: {value}")
            
            print(f"\n2. Testing specific placeholder values...")
            test_placeholders = ['[NAME]', '[DESCRIPTION]', '[TEST_COL]', '<<FIRM_NAME>>', '<<PG_NO>>']
            for placeholder in test_placeholders:
                value = placeholders.get(placeholder, "NOT FOUND")
                print(f"  {placeholder}: {value}")
            
            print(f"\n3. Testing template processor...")
            # Create a test template document
            test_doc = Document()
            paragraph = test_doc.add_paragraph()
            # Add various placeholder types
            paragraph.add_run("Work Name: [NAME]\n")
            paragraph.add_run("Work Description: [DESCRIPTION]\n") 
            paragraph.add_run("Test Column: [TEST_COL]\n")
            paragraph.add_run("Current Date: [CURRENT_DATE]\n")
            paragraph.add_run("Current Time: [CURRENT_TIME]\n")
            paragraph.add_run("Firm Name: <<FIRM_NAME>>\n")
            paragraph.add_run("PG Number: <<PG_NO>>\n")
            paragraph.add_run("User Input: {{TEST_INPUT}}\n")
            
            template_path = "test_investigation_template.docx"
            test_doc.save(template_path)
            
            print(f"Template created with placeholders")
            
            # Test template processor
            processor = TemplateProcessor()
            output_path = "test_investigation_output.docx"
            
            # Test user input data
            user_data = {"TEST_INPUT": "User Input Value"}
            
            print(f"\n4. Processing template...")
            print(f"User data: {user_data}")
            
            # Process the template
            success, message = processor.replace_placeholders(
                template_path, user_data, 1, output_path, firm_placeholders=set()
            )
            
            if success:
                print(f"SUCCESS: Template processing completed - {message}")
                
                # Verify output document content
                output_doc = Document(output_path)
                output_text = ""
                for paragraph in output_doc.paragraphs:
                    output_text += paragraph.text + "\n"
                
                print(f"\n5. Output document analysis:")
                print(f"Raw output text:\n{repr(output_text)}")
                print(f"\nFormatted output text:\n{output_text}")
                
                # Analyze each line
                lines = output_text.strip().split('\n')
                print(f"\n6. Line-by-line analysis:")
                for i, line in enumerate(lines):
                    print(f"  Line {i+1}: {repr(line)}")
                
                # Check for specific replacements with detailed analysis
                print(f"\n7. Replacement verification:")
                expected_replacements = [
                    ("Test Work", "[NAME]", "Work name"),
                    ("Test Description", "[DESCRIPTION]", "Work description"), 
                    ("Test Column Value", "[TEST_COL]", "Test column"),
                    ("Test Firm", "<<FIRM_NAME>>", "Firm name"),
                    ("PG123", "<<PG_NO>>", "PG number"),
                    ("User Input Value", "{{TEST_INPUT}}", "User input")
                ]
                
                for expected_value, placeholder, description in expected_replacements:
                    if expected_value in output_text:
                        print(f"  ✓ {description} ({placeholder}) replaced correctly: '{expected_value}'")
                    else:
                        print(f"  ✗ {description} ({placeholder}) NOT replaced: expected '{expected_value}'")
                        # Check if placeholder still exists
                        if placeholder in output_text:
                            print(f"    → Placeholder '{placeholder}' still present in output")
                        else:
                            print(f"    → Placeholder '{placeholder}' was replaced with something else")
                
                # Check for unreplaced placeholders
                print(f"\n8. Unreplaced placeholder check:")
                unreplaced_placeholders = []
                for placeholder in ['[NAME]', '[DESCRIPTION]', '[TEST_COL]', '[CURRENT_DATE]', '[CURRENT_TIME]',
                                  '<<FIRM_NAME>>', '<<PG_NO>>', '{{TEST_INPUT}}']:
                    if placeholder in output_text:
                        unreplaced_placeholders.append(placeholder)
                
                if unreplaced_placeholders:
                    print(f"  Found unreplaced placeholders: {unreplaced_placeholders}")
                else:
                    print(f"  All placeholders were processed (replaced or transformed)")
                
                # Clean up output file
                if os.path.exists(output_path):
                    os.remove(output_path)
            else:
                print(f"FAILURE: Template processing failed - {message}")
                
            # Clean up template file
            if os.path.exists(template_path):
                os.remove(template_path)
                
        finally:
            # Restore original database path
            config.DATABASE_PATH = original_database_path
            db_utils.DATABASE_PATH = original_database_path
            
    finally:
        # Clean up temporary database
        if os.path.exists(temp_db_path):
            os.unlink(temp_db_path)

def investigate_database_schema():
    print("\n=== DATABASE SCHEMA INVESTIGATION ===\n")
    
    print("1. Current database schema:")
    try:
        from database.managers.database_utils import get_work_columns, get_firm_documents_columns
        
        work_columns = get_work_columns()
        print(f"Works table columns ({len(work_columns)}): {work_columns}")
        
        firm_columns = get_firm_documents_columns()  
        print(f"Firm documents columns ({len(firm_columns)}): {firm_columns}")
        
        print(f"\n2. Expected placeholders based on schema:")
        print("Work placeholders:")
        for col in work_columns:
            print(f"  [{col.upper()}]")
            
        print("Firm placeholders:")
        for col in firm_columns:
            print(f"  <<{col.upper()}>>")
            
    except Exception as e:
        print(f"Error investigating schema: {e}")

def investigate_work_data_provider_static():
    print("\n=== WORK DATA PROVIDER STATIC INVESTIGATION ===\n")
    
    try:
        placeholders = WorkDataProvider.get_available_placeholders_static()
        print(f"Available static placeholders ({len(placeholders)}):")
        
        work_phs = []
        firm_phs = []
        special_phs = []
        
        for key, desc in sorted(placeholders.items()):
            if key.startswith('[') and key.endswith(']'):
                work_phs.append(f"  {key}: {desc}")
            elif key.startswith('<<') and key.endswith('>>'):
                firm_phs.append(f"  {key}: {desc}")
            else:
                special_phs.append(f"  {key}: {desc}")
        
        print(f"\nWork placeholders ({len(work_phs)}):")
        for ph in work_phs:
            print(ph)
            
        print(f"\nFirm placeholders ({len(firm_phs)}):")  
        for ph in firm_phs:
            print(ph)
            
        print(f"\nSpecial placeholders ({len(special_phs)}):")
        for ph in special_phs:
            print(ph)
            
    except Exception as e:
        print(f"Error investigating static placeholders: {e}")

if __name__ == "__main__":
    investigate_database_schema()
    investigate_work_data_provider_static()
    investigate_placeholder_replacement()
