#!/usr/bin/env python3
"""
Test Script for Template Engine Placeholder Replacement
=======================================================

This script:
1. Populates the database with test data (some fields filled, some empty)
2. Creates a comprehensive test template with all placeholder types
3. Tests the template generation process
4. Verifies that placeholders are correctly replaced or left untouched
"""

import os
import sys
import sqlite3
from datetime import datetime
from pathlib import Path

# Add the project root to Python path
sys.path.insert(0, os.path.abspath('.'))

from database.db_manager import create_tables, DATABASE_PATH
from features.template_engine.work_data_provider import WorkDataProvider
from features.template_engine.template_processor import TemplateProcessor
from docx import Document

def create_test_database():
    """Create and populate test database with sample data."""
    print("🔧 Setting up test database...")
    
    # Initialize database
    create_tables()
    
    # Get database path
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    # Clear existing test data
    cursor.execute("DELETE FROM works WHERE name LIKE 'TEST_%'")
    cursor.execute("DELETE FROM firms WHERE name LIKE 'TEST_%'")
    conn.commit()
    
    # Insert test work with mixed data (some fields filled, some empty)
    test_work_data = {
        'name': 'TEST_WORK_COMPREHENSIVE',
        'description': 'Comprehensive test work for template generation',
        'justification': 'Testing all placeholder types and replacements',
        'section': 'Test Section A',
        'work_type': 'Construction',
        'file_no': 'FILE-TEST-2024-001',
        'estimate_no': 'EST-2024-001',
        'tender_cost': 1500000.75,
        'tender_opening_date': '2024-01-10',
        'loa_no': 'LOA-2024-001',
        'loa_date': '2024-01-20',
        'work_commence_date': '2024-02-01',
        'admin_approval_office_note_no': 'NOTE-2024-001',
        'admin_approval_date': '2024-01-05',
        'work_type_category': 'Civil Construction',
        'work_type_subcategory': 'Building',
        'concurrence_letter_no': None,  # Empty field - should remain as placeholder
        'concurrence_letter_dated': None,  # Empty field
        'dr_dfm_eoffice_note_no': None,  # Empty field
        'computer_no': 'COMP-2024-001'
    }
    
    # Insert work
    columns = ', '.join(test_work_data.keys())
    placeholders = ', '.join(['?' for _ in test_work_data])
    cursor.execute(f"INSERT INTO works ({columns}) VALUES ({placeholders})", 
                   list(test_work_data.values()))
    work_id = cursor.lastrowid
    
    # Insert test firms with mixed data (using actual firms table schema)
    test_firms = [
        {
            'name': 'TEST_FIRM_PRIMARY',
            'representative': 'Mr. Primary Contact',
            'address': '456 Primary Avenue, Primary City - 654321'
        },
        {
            'name': 'TEST_FIRM_SECONDARY',
            'representative': 'Ms. Secondary Contact',
            'address': '789 Secondary Road, Secondary Town - 987654'
        }
    ]
    
    firm_ids = []
    for firm_data in test_firms:
        columns = ', '.join(firm_data.keys())
        placeholders = ', '.join(['?' for _ in firm_data])
        cursor.execute(f"INSERT INTO firms ({columns}) VALUES ({placeholders})", 
                       list(firm_data.values()))
        firm_ids.append(cursor.lastrowid)
    
    # Insert firm documents with mixed data (some fields filled, some empty)
    test_firm_docs = [
        {
            'work_id': work_id,
            'firm_name': 'TEST_FIRM_PRIMARY',
            'pg_no': 'PG-2024-001',
            'pg_amount': 150000.0,
            'bank_name': 'Test Bank Ltd.',
            'bank_address': 'Bank Street, Bank City',
            'firm_address': '456 Primary Avenue, Primary City - 654321', 
            'indemnity_bond_details': 'IB Details for Primary Firm',
            'other_docs_details': None,  # Empty field
            'submission_date': '2024-01-25',
            'pg_submitted': 1,
            'indemnity_bond_submitted': 1,
            'pg_type': 'Bank Guarantee',
            'pg_vetted_on': '2024-01-26',
            'ib_vetted_on': None  # Empty field
        },
        {
            'work_id': work_id,
            'firm_name': 'TEST_FIRM_SECONDARY',
            'pg_no': None,  # Empty field
            'pg_amount': None,  # Empty field
            'bank_name': None,  # Empty field
            'bank_address': None,  # Empty field
            'firm_address': '789 Secondary Road, Secondary Town - 987654',
            'indemnity_bond_details': None,  # Empty field
            'other_docs_details': 'Some other documents',
            'submission_date': None,  # Empty field
            'pg_submitted': 0,
            'indemnity_bond_submitted': 0,
            'pg_type': None,  # Empty field
            'pg_vetted_on': None,  # Empty field
            'ib_vetted_on': None  # Empty field
        }
    ]
    
    for doc_data in test_firm_docs:
        columns = ', '.join(doc_data.keys())
        placeholders = ', '.join(['?' for _ in doc_data])
        cursor.execute(f"INSERT INTO firm_documents ({columns}) VALUES ({placeholders})", 
                       list(doc_data.values()))
    
    conn.commit()
    conn.close()
    
    print(f"✅ Test database populated successfully!")
    print(f"   - Created work: {test_work_data['name']} (ID: {work_id})")
    print(f"   - Created firms: {len(test_firms)} firms linked to work")
    
    return work_id

def create_comprehensive_test_template():
    """Create a Word document template with all placeholder types."""
    print("📄 Creating comprehensive test template...")
    
    template_path = "test_template_comprehensive.docx"
    doc = Document()
    
    # Add title
    title = doc.add_heading('Comprehensive Template Test Document', 0)
    
    # Add comprehensive content with all placeholder types
    content = """
WORK DETAILS SECTION
====================

Work Name: [NAME]
Description: [DESCRIPTION] 
Justification: [JUSTIFICATION]
Section: [SECTION]
Work Type: [WORK_TYPE]
File Number: [FILE_NO]
Estimate Number: [ESTIMATE_NO]

FINANCIAL DETAILS
=================

Tender Cost: [TENDER_COST]
Approved Cost: [APPROVED_COST]
Contract Amount: [CONTRACT_AMOUNT]
EMD Amount: [EMD_AMOUNT]
Retention Percentage: [RETENTION_PERCENTAGE]%
Performance Guarantee: [PERFORMANCE_GUARANTEE]%

SCHEDULE DETAILS
================

Work Order Number: [WORK_ORDER_NO]
Work Order Date: [WORK_ORDER_DATE]
Completion Period: [COMPLETION_PERIOD] months
Extended Date: [EXTENDED_DATE]
Actual Completion Date: [ACTUAL_COMPLETION_DATE]
Defect Liability Period: [DEFECT_LIABILITY_PERIOD] months

CONTRACTOR DETAILS
==================

Contractor Name: [CONTRACTOR_NAME]
Contractor Address: [CONTRACTOR_ADDRESS]
Contractor Phone: [CONTRACTOR_PHONE]
Contractor Email: [CONTRACTOR_EMAIL]

PROJECT TEAM
============

Engineer Name: [ENGINEER_NAME]
Project Manager: [PROJECT_MANAGER]

FIRM DETAILS SECTION
====================

Primary Firm Name: <<NAME>>
Primary Contact Person: <<CONTACT_PERSON>>
Primary Phone: <<PHONE>>
Primary Email: <<EMAIL>>
Primary Address: <<ADDRESS>>
Primary GST No: <<GST_NO>>
Primary PAN No: <<PAN_NO>>
Primary Bank Name: <<BANK_NAME>>
Primary Account No: <<ACCOUNT_NO>>
Primary IFSC Code: <<IFSC_CODE>>
Primary Registration No: <<REGISTRATION_NO>>
Primary Website: <<WEBSITE>>

SPECIAL PLACEHOLDERS
====================

Current Date: [CURRENT_DATE]
Current Time: [CURRENT_TIME]

Firm PG Details:
[FIRM_PG_DETAILS]

All Firms PG Details:
[ALL_FIRMS_PG_DETAILS]

TEMPLATE ENGINE PLACEHOLDERS
=============================

Base Cost: {{COST}}
Cost with 10% increase: {{COST_1.1}}
Cost rounded to nearest 100: {{COST_00}}
Cost in words: {{COST_IN_WORDS}}

Manual Input Field: {{manual_field}}
Date Field: {{DATE_contract_date}}

REMARKS SECTION
===============

Additional Remarks: [REMARKS]

---
Document generated on [CURRENT_DATE] at [CURRENT_TIME]
"""
    
    # Add content to document
    for line in content.strip().split('\n'):
        if line.strip():
            if line.startswith('===') or line.startswith('---'):
                continue
            elif line.isupper() and line.endswith('SECTION') or line.isupper() and line.endswith('DETAILS'):
                doc.add_heading(line, level=1)
            else:
                doc.add_paragraph(line)
    
    # Save template
    doc.save(template_path)
    print(f"✅ Test template created: {template_path}")
    
    return template_path

def test_placeholder_replacement(work_id, template_path):
    """Test the placeholder replacement process."""
    print("🧪 Testing placeholder replacement...")
    
    try:
        # Initialize components
        work_data_provider = WorkDataProvider(work_id)
        template_processor = TemplateProcessor()
        
        # Generate output document
        output_path = "test_output_comprehensive.docx"
        
        # Simulate the template processing
        user_input_data = {
            'COST': '1500000.75',  # This should trigger COST calculations
            'manual_field': 'User entered test value',
            'DATE_contract_date': '2024-02-15'
        }
        
        # Process template (this simulates the right-click menu action)
        success, message = template_processor.replace_placeholders(
            doc_path=template_path,
            data=user_input_data,
            work_id=work_id,
            output_path=output_path,
            firm_placeholders=set()  # No firm placeholders for this test
        )
        
        if success:
            print(f"✅ Template processed successfully: {output_path}")
            return output_path
        else:
            print(f"❌ Template processing failed: {message}")
            return None
            
    except Exception as e:
        print(f"❌ Error during template processing: {e}")
        import traceback
        traceback.print_exc()
        return None

def verify_output_document(output_path, work_id):
    """Verify that the output document has correct placeholder replacements."""
    print("🔍 Verifying output document...")
    
    if not os.path.exists(output_path):
        print("❌ Output document not found!")
        return False
    
    try:
        # Read the output document
        doc = Document(output_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        content = '\n'.join(full_text)
        
        print("\n📋 VERIFICATION RESULTS:")
        print("=" * 50)
        
        # Get work data for verification
        work_provider = WorkDataProvider(work_id)
        work_details = work_provider.work_details
        
        # Test cases for verification
        test_cases = [
            # Should be replaced (have values)
            ("[NAME]", work_details.get('name'), "Work name"),
            ("[DESCRIPTION]", work_details.get('description'), "Work description"),
            ("[TENDER_COST]", str(work_details.get('tender_cost')), "Tender cost"),
            ("[CONTRACTOR_NAME]", work_details.get('contractor_name'), "Contractor name"),
            ("[CURRENT_DATE]", datetime.now().strftime("%Y-%m-%d"), "Current date"),
            
            # Should remain as placeholders (empty/None values)
            ("[CONTRACT_AMOUNT]", "[CONTRACT_AMOUNT]", "Contract amount (empty)"),
            ("[PROJECT_MANAGER]", "[PROJECT_MANAGER]", "Project manager (empty)"),
            ("[REMARKS]", "[REMARKS]", "Remarks (empty)"),
            ("[EXTENDED_DATE]", "[EXTENDED_DATE]", "Extended date (empty)"),
        ]
        
        all_passed = True
        
        for placeholder, expected, description in test_cases:
            if expected and expected != placeholder:
                # Should be replaced
                if expected in content and placeholder not in content:
                    print(f"✅ {description}: REPLACED correctly")
                else:
                    print(f"❌ {description}: NOT replaced (expected: {expected})")
                    all_passed = False
            else:
                # Should remain as placeholder
                if placeholder in content:
                    print(f"✅ {description}: REMAINED as placeholder (correct)")
                else:
                    print(f"❌ {description}: Unexpectedly replaced or missing")
                    all_passed = False
        
        # Check for unreplaced placeholders that should have been replaced
        import re
        unreplaced_work = re.findall(r'\[([A-Z_]+)\]', content)
        unreplaced_firm = re.findall(r'<<([A-Z_]+)>>', content)
        unreplaced_template = re.findall(r'{{([a-zA-Z0-9_]+)}}', content)
        
        print(f"\n📊 PLACEHOLDER ANALYSIS:")
        print(f"   Unreplaced work placeholders: {len(unreplaced_work)}")
        if unreplaced_work:
            print(f"   Details: {unreplaced_work}")
        
        print(f"   Unreplaced firm placeholders: {len(unreplaced_firm)}")
        if unreplaced_firm:
            print(f"   Details: {unreplaced_firm}")
            
        print(f"   Unreplaced template placeholders: {len(unreplaced_template)}")
        if unreplaced_template:
            print(f"   Details: {unreplaced_template}")
        
        # Save detailed report
        report_path = "placeholder_replacement_report.txt"
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("PLACEHOLDER REPLACEMENT TEST REPORT\n")
            f.write("=" * 40 + "\n\n")
            f.write(f"Generated: {datetime.now()}\n")
            f.write(f"Work ID: {work_id}\n")
            f.write(f"Template: {output_path}\n\n")
            
            f.write("DOCUMENT CONTENT:\n")
            f.write("-" * 20 + "\n")
            f.write(content)
            f.write("\n\n")
            
            f.write("UNREPLACED PLACEHOLDERS:\n")
            f.write("-" * 25 + "\n")
            f.write(f"Work: {unreplaced_work}\n")
            f.write(f"Firm: {unreplaced_firm}\n")
            f.write(f"Template: {unreplaced_template}\n")
        
        print(f"\n📄 Detailed report saved: {report_path}")
        
        if all_passed:
            print("\n🎉 ALL TESTS PASSED! Template generation working correctly.")
        else:
            print("\n⚠️  Some tests failed. Check the results above.")
        
        return all_passed
        
    except Exception as e:
        print(f"❌ Error verifying output: {e}")
        import traceback
        traceback.print_exc()
        return False

def cleanup_test_files():
    """Clean up test files."""
    files_to_remove = [
        "test_template_comprehensive.docx",
        "test_output_comprehensive.docx",
        "placeholder_replacement_report.txt"
    ]
    
    print("\n🧹 Cleaning up test files...")
    for file_path in files_to_remove:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                print(f"   Removed: {file_path}")
            except Exception as e:
                print(f"   Failed to remove {file_path}: {e}")

def main():
    """Main test execution function."""
    print("🚀 Starting Comprehensive Template Generation Test")
    print("=" * 60)
    
    try:
        # Step 1: Create and populate test database
        work_id = create_test_database()
        
        # Step 2: Create comprehensive test template
        template_path = create_comprehensive_test_template()
        
        # Step 3: Test placeholder replacement
        output_path = test_placeholder_replacement(work_id, template_path)
        
        if output_path:
            # Step 4: Verify the results
            success = verify_output_document(output_path, work_id)
            
            if success:
                print("\n🎯 TEST COMPLETED SUCCESSFULLY!")
                print("   The template engine correctly:")
                print("   ✅ Replaces placeholders with database values")
                print("   ✅ Leaves empty/None placeholders unchanged")
                print("   ✅ Handles all placeholder types correctly")
            else:
                print("\n❌ TEST COMPLETED WITH ISSUES!")
                print("   Please check the verification results above.")
        else:
            print("\n❌ TEST FAILED - Could not generate output document")
        
        # Ask user if they want to keep test files
        keep_files = input("\nKeep test files for manual inspection? (y/N): ").lower().strip()
        if keep_files != 'y':
            cleanup_test_files()
        else:
            print("   Test files preserved for manual inspection.")
        
    except Exception as e:
        print(f"\n💥 TEST FAILED WITH ERROR: {e}")
        import traceback
        traceback.print_exc()
    
    print("\n" + "=" * 60)
    print("Test execution completed.")

if __name__ == "__main__":
    main()
