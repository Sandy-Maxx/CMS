#!/usr/bin/env python3

import sys
import os
import logging
from docx import Document

# Configure logging to see debug output
logging.basicConfig(
    level=logging.DEBUG,
    format='%(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)

# Add current directory to Python path
sys.path.append('.')

from features.AutodocGen.data_fetcher import DataFetcher
from features.AutodocGen.document_generator import DocumentGenerator
from config import DATABASE_PATH

def test_placeholder_replacement():
    """Test placeholder replacement with logging to see what's happening"""
    
    print("Starting placeholder replacement test...")
    
    # Initialize the data fetcher and document generator
    try:
        data_fetcher = DataFetcher(DATABASE_PATH)
        generator = DocumentGenerator(data_fetcher)
        
        # Test data - use an existing work ID from the database
        test_data = {
            'work_id': 1,  # Assuming work_id 1 exists
            'firm_name': 'Test Firm'
        }
        
        # Template and output paths
        template_path = 'Templates/Letters/letter for IB, PG.docx'
        output_path = 'debug_output.docx'
        
        print(f"Using template: {template_path}")
        print(f"Output will be saved to: {output_path}")
        print(f"Test data: {test_data}")
        
        # Generate the document
        generator.generate(template_path, test_data, output_path, is_firm_specific=True)
        
        print("Document generation completed!")
        
        # Read the output to see what happened to the placeholders
        print("\n--- OUTPUT DOCUMENT CONTENT ---")
        try:
            output_doc = Document(output_path)
            for i, paragraph in enumerate(output_doc.paragraphs):
                text = paragraph.text
                if text.strip():
                    print(f"Paragraph {i}: {repr(text)}")
                    
            # Check tables too
            for table_idx, table in enumerate(output_doc.tables):
                print(f"Table {table_idx}:")
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text
                        if cell_text.strip():
                            print(f"  Row {row_idx}, Cell {cell_idx}: {repr(cell_text)}")
        except Exception as e:
            print(f"Error reading output document: {e}")
        
    except Exception as e:
        print(f"Error during test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_placeholder_replacement()
