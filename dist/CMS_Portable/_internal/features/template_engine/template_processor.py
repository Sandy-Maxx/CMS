from docx import Document
import re
from features.template_engine.special_placeholder_handler import evaluate_special_placeholder
from .work_data_provider import WorkDataProvider

class TemplateProcessor:
    def __init__(self):
        pass

    def extract_placeholders(self, doc_path):
        document = Document(doc_path)
        user_input_placeholders = set()
        work_level_placeholders = set()
        firm_level_placeholders = set()

        def _extract_from_text(text_content):
            # Find all user input placeholders: {{placeholder}}
            found_user_input = re.findall(r"\{\{([a-zA-Z0-9_.]+)\}\}", text_content)
            for p in found_user_input:
                user_input_placeholders.add(p)
            
            # Find all work-level placeholders: [PLACEHOLDER]
            found_work_level = re.findall(r"\[([A-Za-z0-9_]+)\]", text_content)
            for p in found_work_level:
                work_level_placeholders.add(p)
            
            # Find all firm-level placeholders: <<PLACEHOLDER>>
            found_firm_level = re.findall(r"<<([A-Za-z0-9_]+)>>", text_content)
            for p in found_firm_level:
                firm_level_placeholders.add(p)
            
        # Extract from paragraphs in the main body
        for paragraph in document.paragraphs:
            _extract_from_text(paragraph.text)

        # Extract from tables in the main body
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _extract_from_text(cell.text)

        # Extract from headers and footers
        for section in document.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                _extract_from_text(paragraph.text)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(cell.text)

            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                _extract_from_text(paragraph.text)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(cell.text)

        # Filter for base placeholders only (for user input placeholders)
        base_placeholders_for_gui = set()
        for p_name in user_input_placeholders:
            is_derived = False
            if p_name.startswith(("COST", "COSTAMC", "COSTRP", "COSTCON", "COSTCAMC")):
                if re.search(r"_([0-9.]+)$|_IN_WORDS$|_0$|_00$", p_name):
                    is_derived = True
            
            if p_name.startswith("DATE"):
                is_derived = False

            if not is_derived:
                base_placeholders_for_gui.add(p_name)
        
        # For backward compatibility, return just the user input placeholders by default
        # New method extract_all_placeholders() can be used to get the full dictionary
        return base_placeholders_for_gui
    
    def extract_all_placeholders(self, doc_path):
        """
        Enhanced version that returns all placeholder types for debugging and advanced use.
        Returns a dictionary with different placeholder types.
        """
        document = Document(doc_path)
        user_input_placeholders = set()
        work_level_placeholders = set()
        firm_level_placeholders = set()

        def _extract_from_text(text_content):
            # Find all user input placeholders: {{placeholder}}
            found_user_input = re.findall(r"\{\{([a-zA-Z0-9_.]+)\}\}", text_content)
            for p in found_user_input:
                user_input_placeholders.add(p)
            
            # Find all work-level placeholders: [PLACEHOLDER]
            found_work_level = re.findall(r"\[([A-Za-z0-9_]+)\]", text_content)
            for p in found_work_level:
                work_level_placeholders.add(p)
            
            # Find all firm-level placeholders: <<PLACEHOLDER>>
            found_firm_level = re.findall(r"<<([A-Za-z0-9_]+)>>", text_content)
            for p in found_firm_level:
                firm_level_placeholders.add(p)
            
        # Extract from paragraphs in the main body
        for paragraph in document.paragraphs:
            _extract_from_text(paragraph.text)

        # Extract from tables in the main body
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _extract_from_text(cell.text)

        # Extract from headers and footers
        for section in document.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                _extract_from_text(paragraph.text)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(cell.text)

            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                _extract_from_text(paragraph.text)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(cell.text)

        # Filter for base placeholders only (for user input placeholders)
        base_placeholders_for_gui = set()
        for p_name in user_input_placeholders:
            is_derived = False
            if p_name.startswith(("COST", "COSTAMC", "COSTRP", "COSTCON", "COSTCAMC")):
                if re.search(r"_([0-9.]+)$|_IN_WORDS$|_0$|_00$", p_name):
                    is_derived = True
            
            if p_name.startswith("DATE"):
                is_derived = False

            if not is_derived:
                base_placeholders_for_gui.add(p_name)
        
        # Return a dictionary with all placeholder types for advanced debugging
        return {
            'user_input': base_placeholders_for_gui,
            'work_level': work_level_placeholders,
            'firm_level': firm_level_placeholders,
            'all_user_input': user_input_placeholders
        }

    def generate_letters_for_firms(self, doc_path, data, work_id, output_path):
        work_data_provider = WorkDataProvider(work_id)
        firm_names = work_data_provider.get_firm_names_list().split(', ')
        if not firm_names:
            return False, "No firms found for this work."

        master_document = Document()
        for i, firm_name in enumerate(firm_names):
            # Create a copy of the template for each firm
            firm_document = Document(doc_path)
            firm_data = data.copy()
            firm_data['firm_name'] = firm_name # Add firm name to data
            
            # Get firm-specific placeholder data
            firm_specific_work_data_provider = WorkDataProvider(work_id)
            firm_placeholder_data = firm_specific_work_data_provider.generate_placeholders()
            
            # Update firm-specific data if needed
            firm_doc_data = firm_specific_work_data_provider.get_firm_document_data(firm_name)
            if firm_doc_data:
                # Update firm-level placeholders with specific firm's data
                firm_columns = ['firm_name', 'pg_submitted', 'pg_no', 'submission_date', 'pg_amount', 'bank_name', 'bank_address']
                for column in firm_columns:
                    if column in firm_doc_data:
                        firm_placeholder_data[f'<<{column.upper()}>>'] = firm_doc_data[column]

            # Process all placeholders for this firm using two-pass system
            for paragraph in firm_document.paragraphs:
                self._replace_placeholders_two_pass(paragraph, firm_data, firm_placeholder_data)
            for table in firm_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_two_pass(paragraph, firm_data, firm_placeholder_data)
            
            # Process headers and footers for this firm
            for section in firm_document.sections:
                # Headers
                header = section.header
                for paragraph in header.paragraphs:
                    self._replace_placeholders_two_pass(paragraph, firm_data, firm_placeholder_data)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_placeholders_two_pass(paragraph, firm_data, firm_placeholder_data)

                # Footers
                footer = section.footer
                for paragraph in footer.paragraphs:
                    self._replace_placeholders_two_pass(paragraph, firm_data, firm_placeholder_data)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_placeholders_two_pass(paragraph, firm_data, firm_placeholder_data)

            # Add the processed letter to the master document
            for element in firm_document.element.body:
                master_document.element.body.append(element)
            
            # Add a page break after each letter except the last one
            if i < len(firm_names) - 1:
                master_document.add_page_break()

        master_document.save(output_path)
        return True, "Letters generated successfully."

    def replace_placeholders(self, doc_path, data, work_id, output_path, firm_placeholders):
        if firm_placeholders:
            return self.generate_letters_for_firms(doc_path, data, work_id, output_path)

        document = Document(doc_path)
        work_data_provider = WorkDataProvider(work_id)
        
        # Get consolidated placeholder data for two-pass replacement
        placeholder_data = work_data_provider.generate_placeholders()
        
        # Two-pass replacement system:
        # Pass 1: Replace work-level [PLACEHOLDER] tokens
        # Pass 2: Replace firm-level <<PLACEHOLDER>> tokens
        
        # Process paragraphs in the main body
        for paragraph in document.paragraphs:
            self._replace_placeholders_two_pass(paragraph, data, placeholder_data)

        # Process tables in the main body
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_two_pass(paragraph, data, placeholder_data)

        # Process headers and footers
        for section in document.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_placeholders_two_pass(paragraph, data, placeholder_data)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_two_pass(paragraph, data, placeholder_data)

            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_placeholders_two_pass(paragraph, data, placeholder_data)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_two_pass(paragraph, data, placeholder_data)

        document.save(output_path)
        return True, "Document generated successfully."

    def _replace_in_paragraph(self, paragraph, data, work_data_provider, firm_name=None):
        full_text = "".join([run.text for run in paragraph.runs])
        
        # Regex for user input placeholders only
        matches = list(re.finditer(r"\{\{([a-zA-Z0-9_.]+)\}\}", full_text))

        if not matches:
            return

        new_runs_data = []
        last_idx = 0

        for match in matches:
            placeholder_full = match.group(0)
            user_input_ph = match.group(1) # Correctly access the captured group

            if match.start() > last_idx:
                new_runs_data.append({'text': full_text[last_idx:match.start()], 'style': None})

            replacement_value = evaluate_special_placeholder(user_input_ph, data)
            if replacement_value is None or str(replacement_value).strip() == "":
                new_runs_data.append({'text': placeholder_full, 'style': None})
            else:
                new_runs_data.append({'text': str(replacement_value), 'style': None})

            last_idx = match.end()

        if last_idx < len(full_text):
            new_runs_data.append({'text': full_text[last_idx:], 'style': None})

        # Clear existing runs and add new ones
        for i in reversed(range(len(paragraph.runs))):
            p = paragraph.runs[i]
            p.element.getparent().remove(p.element)

        for run_data in new_runs_data:
            new_run = paragraph.add_run(run_data['text'])
            if run_data['style']:
                new_run.style = run_data['style']

    def _replace_placeholders_two_pass(self, paragraph, user_data, placeholder_data):
        """
        Enhanced merge routine that performs two passes:
        1. Replace work-level [PLACEHOLDER] tokens
        2. Replace firm-level <<PLACEHOLDER>> tokens
        Fallback: unknown placeholders remain untouched for easier debugging
        """
        full_text = "".join([run.text for run in paragraph.runs])
        
        # Pass 1: Replace work-level [PLACEHOLDER] tokens
        full_text = self._replace_work_level_placeholders(full_text, user_data, placeholder_data)
        
        # Pass 2: Replace firm-level <<PLACEHOLDER>> tokens  
        full_text = self._replace_firm_level_placeholders(full_text, placeholder_data)
        
        # Pass 3: Replace user input {{placeholder}} tokens (existing functionality)
        full_text = self._replace_user_input_placeholders(full_text, user_data)
        
        # Update paragraph with the processed text
        self._update_paragraph_text(paragraph, full_text)
    
    def _replace_work_level_placeholders(self, text, user_data, placeholder_data):
        """
        Pass 1: Replace work-level [PLACEHOLDER] tokens
        """
        # Find all work-level placeholders: [PLACEHOLDER]
        matches = list(re.finditer(r'\[([A-Za-z0-9_]+)\]', text))
        
        if not matches:
            return text
            
        result = text
        # Process matches in reverse order to avoid index shifting
        for match in reversed(matches):
            placeholder_full = match.group(0)  # [PLACEHOLDER]
            placeholder_key = match.group(1)   # PLACEHOLDER
            
            # Try to find replacement value
            replacement_value = None
            
            # Check in placeholder_data first (generated from WorkDataProvider)
            if placeholder_full in placeholder_data:
                replacement_value = placeholder_data[placeholder_full]
            # Check in user_data as fallback
            elif placeholder_key in user_data:
                replacement_value = user_data[placeholder_key]
            elif placeholder_key.lower() in user_data:
                replacement_value = user_data[placeholder_key.lower()]
            
            # Replace if value found, otherwise leave unchanged for debugging
            if replacement_value is not None and str(replacement_value).strip() != "":
                result = result[:match.start()] + str(replacement_value) + result[match.end():]
            # If no replacement found, leave the placeholder untouched for debugging
            
        return result
    
    def _replace_firm_level_placeholders(self, text, placeholder_data):
        """
        Pass 2: Replace firm-level <<PLACEHOLDER>> tokens
        This may loop over firms or pick a default firm based on placeholder_data
        """
        # Find all firm-level placeholders: <<PLACEHOLDER>>
        matches = list(re.finditer(r'<<([A-Za-z0-9_]+)>>', text))
        
        if not matches:
            return text
            
        result = text
        # Process matches in reverse order to avoid index shifting
        for match in reversed(matches):
            placeholder_full = match.group(0)  # <<PLACEHOLDER>>
            placeholder_key = match.group(1)   # PLACEHOLDER
            
            # Try to find replacement value
            replacement_value = None
            
            # Check in placeholder_data (generated from WorkDataProvider)
            if placeholder_full in placeholder_data:
                replacement_value = placeholder_data[placeholder_full]
            # Try with different case variations
            elif f'<<{placeholder_key.upper()}>>' in placeholder_data:
                replacement_value = placeholder_data[f'<<{placeholder_key.upper()}>>']  
            elif f'<<{placeholder_key.lower()}>>' in placeholder_data:
                replacement_value = placeholder_data[f'<<{placeholder_key.lower()}>>']    
            
            # Replace if value found, otherwise leave unchanged for debugging
            if replacement_value is not None and str(replacement_value).strip() != "":
                result = result[:match.start()] + str(replacement_value) + result[match.end():]
            # If no replacement found, leave the placeholder untouched for debugging
            
        return result
    
    def _replace_user_input_placeholders(self, text, user_data):
        """
        Pass 3: Replace user input {{placeholder}} tokens (existing functionality)
        """
        # Find all user input placeholders: {{placeholder}}
        matches = list(re.finditer(r'\{\{([a-zA-Z0-9_.]+)\}\}', text))
        
        if not matches:
            return text
            
        result = text
        # Process matches in reverse order to avoid index shifting
        for match in reversed(matches):
            placeholder_full = match.group(0)  # {{placeholder}}
            placeholder_key = match.group(1)   # placeholder
            
            # Use existing special placeholder handler
            replacement_value = evaluate_special_placeholder(placeholder_key, user_data)
            
            # Replace if value found and not empty, otherwise leave unchanged for debugging
            if replacement_value is not None and str(replacement_value).strip() != "":
                # Check if it's returning the original placeholder (no replacement found)
                if str(replacement_value) != placeholder_full:
                    result = result[:match.start()] + str(replacement_value) + result[match.end():]
            # If no replacement found, leave the placeholder untouched for debugging
            
        return result
    
    def _update_paragraph_text(self, paragraph, new_text):
        """
        Update paragraph with new text while preserving formatting where possible
        """
        # Clear existing runs
        for i in reversed(range(len(paragraph.runs))):
            p = paragraph.runs[i]
            p.element.getparent().remove(p.element)
        
        # Add the new text as a single run
        if new_text:
            paragraph.add_run(new_text)
