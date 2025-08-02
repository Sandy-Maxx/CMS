import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from features.AutodocGen.constants import LETTERS_TEMPLATE_DIR, OFFICE_NOTES_TEMPLATE_DIR, USER_PLACEHOLDER_PATTERN, WORK_DATA_PLACEHOLDER_PATTERN, FIRM_PLACEHOLDER_PATTERN
from features.AutodocGen.placeholder_parser import PlaceholderParser
from features.AutodocGen.data_fetcher import DataFetcher
from features.AutodocGen.document_generator import DocumentGenerator
from features.AutodocGen.firm_selector_dialog import FirmSelectorDialog
from utils.helpers import show_toast

class AutodocManager:
    def __init__(self, master, db_path):
        self.master = master
        self.db_path = db_path
        self.placeholder_parser = PlaceholderParser()
        self.data_fetcher = DataFetcher(db_path)
        self.document_generator = DocumentGenerator(self.data_fetcher)

    def generate_document(self, work_id, doc_type):
        template_dir = ""
        if doc_type == "Letters":
            template_dir = LETTERS_TEMPLATE_DIR
        elif doc_type == "OfficeNotes":
            template_dir = OFFICE_NOTES_TEMPLATE_DIR
        else:
            show_toast(self.master, "Invalid document type.", "error")
            return

        # 1. Select Template
        template_path = self._select_template(template_dir)
        if not template_path:
            show_toast(self.master, "Template selection cancelled.", "info")
            return

        # 2. Extract Placeholders
        try:
            user_placeholders, work_data_placeholders, firm_placeholders, all_firms_pg_details_placeholders = self.placeholder_parser.extract_placeholders(template_path)
        except Exception as e:
            show_toast(self.master, f"Error extracting placeholders: {e}", "error")
            return

        # 3. Fetch Dynamic Data
        dynamic_data = {}
        try:
            # Fetch work-related data
            work_data = self.data_fetcher.fetch_work_data(work_id)
            if work_data:
                dynamic_data.update(work_data)
            
            # Fetch user-input data (if any, from a saved source or prompt)
            # For now, we'll assume user input placeholders are filled manually or from a simple source
            # In a real scenario, you might have a UI to collect these or load from a config
            for ph in user_placeholders:
                # This is a simplification. In a full implementation, you'd prompt the user
                # or load from a saved configuration for these values.
                dynamic_data[ph] = f"<USER_INPUT_{ph}>" 

        except Exception as e:
            show_toast(self.master, f"Error fetching dynamic data: {e}", "error")
            return

        # 4. Check if template needs multi-firm generation
        is_multi_firm_template = self._is_multi_firm_template(template_path)
        
        # 5. Handle Firm Selection (if firm-specific placeholders exist and not multi-firm)
        selected_firm_name = None
        if firm_placeholders and not is_multi_firm_template:
            firms_for_work = self.data_fetcher.fetch_firms_for_work(work_id)
            if not firms_for_work:
                show_toast(self.master, "No firms found for this work to generate firm-specific documents.", "warning")
                return
            
            dialog = FirmSelectorDialog(self.master, firms_for_work)
            self.master.wait_window(dialog)
            selected_firm_name = dialog.selected_firm

            if not selected_firm_name:
                show_toast(self.master, "Firm selection cancelled.", "info")
                return
            
            # Fetch firm-specific data
            firm_data = self.data_fetcher.fetch_firm_data(selected_firm_name, work_id)
            if firm_data:
                dynamic_data.update(firm_data)
            dynamic_data['firm_name'] = selected_firm_name # Ensure firm_name is available

        # 6. Generate Document
        try:
            output_file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                initialfile=f"Generated_{os.path.basename(template_path)}"
            )
            if not output_file_path:
                show_toast(self.master, "Document save cancelled.", "info")
                return

            if is_multi_firm_template:
                # Generate documents for all firms
                self._generate_multi_firm_document(template_path, dynamic_data, output_file_path, work_id)
            else:
                # Generate single document
                self.document_generator.generate(template_path, dynamic_data, output_file_path, selected_firm_name is not None)
                show_toast(self.master, f"Document generated successfully to {output_file_path}", "success")
        except Exception as e:
            show_toast(self.master, f"Error generating document: {e}", "error")

    def _is_multi_firm_template(self, template_path):
        """Check if the template has a specific identifier for multi-firm generation."""
        try:
            document = Document(template_path)
            for paragraph in document.paragraphs:
                if "MULTI_FIRM_IDENTIFIER" in paragraph.text:
                    return True
        except Exception as e:
            print(f"Error reading template for multi-firm identifier: {e}")
        return False

    def _generate_multi_firm_document(self, template_path, dynamic_data, output_file_path, work_id):
        """Generate a combined document with letters for all firms."""
        firms_for_work = self.data_fetcher.fetch_firms_for_work(work_id)
        if not firms_for_work:
            show_toast(self.master, "No firms found for this work.", "warning")
            return

        # Generate first firm document as base
        first_firm = firms_for_work[0]
        firm_data = self.data_fetcher.fetch_firm_data(first_firm, work_id)
        if firm_data:
            dynamic_data.update(firm_data)
        dynamic_data['firm_name'] = first_firm
        
        # Remove the identifier from the template before processing
        self._remove_multi_firm_identifier(template_path)
        
        # Generate the base document with first firm
        self.document_generator.generate(template_path, dynamic_data, output_file_path, is_firm_specific=True)
        base_document = Document(output_file_path)
        
        # Generate and append documents for remaining firms
        for firm_name in firms_for_work[1:]:
            firm_data = self.data_fetcher.fetch_firm_data(firm_name, work_id)
            if firm_data:
                dynamic_data.update(firm_data)
            dynamic_data['firm_name'] = firm_name

            temp_output_path = output_file_path.replace(".docx", f"_temp_{firm_name}.docx")
            self.document_generator.generate(template_path, dynamic_data, temp_output_path, is_firm_specific=True)

            # Add page break before next firm's letter
            from docx.shared import Inches
            from docx.enum.text import WD_BREAK
            base_document.add_page_break()
            
            # Load and append firm document content
            firm_document = Document(temp_output_path)
            for element in firm_document._body._element:
                base_document._body._element.append(element)
            
            # Clean up temp file
            try:
                os.remove(temp_output_path)
            except:
                pass

        base_document.save(output_file_path)
        show_toast(self.master, f"Document for all {len(firms_for_work)} firms generated successfully to {output_file_path}", "success")
        
    def _remove_multi_firm_identifier(self, template_path):
        """Remove the multi-firm identifier from template before processing."""
        try:
            document = Document(template_path)
            for paragraph in document.paragraphs:
                if "MULTI_FIRM_IDENTIFIER" in paragraph.text:
                    paragraph.text = paragraph.text.replace("MULTI_FIRM_IDENTIFIER", "")
            # Save the cleaned template temporarily
            document.save(template_path)
        except Exception as e:
            print(f"Warning: Could not remove multi-firm identifier: {e}")

        

    def _select_template(self, template_dir):
        if not os.path.exists(template_dir):
            os.makedirs(template_dir) # Create directory if it doesn't exist
            show_toast(self.master, f"Template directory created: {template_dir}", "info")

        file_path = filedialog.askopenfilename(
            initialdir=template_dir,
            title="Select Template",
            filetypes=[("Word Documents", "*.docx")]
        )
        return file_path
