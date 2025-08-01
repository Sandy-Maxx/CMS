from docx import Document
import re
from datetime import datetime
from features.AutodocGen.constants import USER_PLACEHOLDER_PATTERN, WORK_DATA_PLACEHOLDER_PATTERN, FIRM_PLACEHOLDER_PATTERN, ALL_FIRMS_PG_DETAILS_PATTERN
from features.template_engine.special_placeholder_handler import evaluate_special_placeholder
from features.template_engine.work_data_provider import WorkDataProvider
from features.AutodocGen.pg_details_formatter import PGDetailsFormatter
from features.AutodocGen.enquiry_table_formatter import EnquiryTableFormatter
from database.db_manager import get_unique_firm_names_by_work_id
from utils.helpers import format_currency_inr

class DocumentGenerator:
    def __init__(self, data_fetcher):
        self.data_fetcher = data_fetcher
        self.pg_formatter = PGDetailsFormatter()
        self.enquiry_table_formatter = EnquiryTableFormatter()

    def generate(self, template_path, data, output_path, is_firm_specific=False):
        document = Document(template_path)
        work_id = data.get('work_id')
        work_data_provider = WorkDataProvider(work_id)
        
        # Initialize table insertions list
        self._table_insertions = []

        # Process all paragraphs in the document
        for paragraph in document.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, data, work_data_provider, is_firm_specific)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, data, work_data_provider, is_firm_specific)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, data, work_data_provider, is_firm_specific)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(paragraph, data, work_data_provider, is_firm_specific)
            
            for paragraph in section.footer.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, data, work_data_provider, is_firm_specific)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(paragraph, data, work_data_provider, is_firm_specific)

        # Process table insertions
        self._process_table_insertions(document)

        document.save(output_path)

    def _replace_placeholders_in_paragraph(self, paragraph, data, work_data_provider, is_firm_specific):
        # Regex to find [PLACEHOLDER:FORMAT] or <<PLACEHOLDER>>
        placeholder_regex = r"(\[([\w_:-]+)\])|(<<([\w_]+)>>)"
        
        full_text = "".join(run.text for run in paragraph.runs)

        # Create a dictionary to hold all replacements
        replacements = {}

        # Handle [ALL_FIRMS_PG_DETAILS] separately
        if "[ALL_FIRMS_PG_DETAILS]" in full_text:
            work_id = data.get('work_id')
            if work_id:
                pg_details = self.data_fetcher.fetch_all_firms_pg_details(work_id)
                replacement_value = self.pg_formatter.format_pg_details(pg_details, work_id)
            else:
                replacement_value = "N/A (Work ID not available)"
            replacements["[ALL_FIRMS_PG_DETAILS]"] = str(replacement_value)
        
        # Handle [ENQUIRY_TABLE] separately - this creates a table
        if "[ENQUIRY_TABLE]" in full_text:
            work_id = data.get('work_id')
            if work_id:
                # Get the first firm as reference firm for ELS KYN Estimate
                firm_names = get_unique_firm_names_by_work_id(work_id)
                reference_firm = firm_names[0] if firm_names else None
                
                # Replace the placeholder text with a marker for table insertion
                replacements["[ENQUIRY_TABLE]"] = "<<TABLE_INSERT_MARKER>>"
                
                # Store table data for later insertion
                if not hasattr(self, '_table_insertions'):
                    self._table_insertions = []
                self._table_insertions.append({
                    'paragraph': paragraph,
                    'work_id': work_id,
                    'reference_firm': reference_firm
                })

        # Find and process all other placeholders
        for match in re.finditer(placeholder_regex, full_text):
            placeholder_full = match.group(0)
            
            if placeholder_full in replacements:
                continue

            key = match.group(2) or match.group(4)
            replacement_value = None

            # Handle [PLACEHOLDER]
            if match.group(1):
                if key.upper().startswith("DATE:"):
                    try:
                        date_format = key.split(":")[1]
                        py_format = date_format.replace("DD", "%d").replace("MM", "%m").replace("YYYY", "%Y")
                        replacement_value = datetime.now().strftime(py_format)
                    except Exception:
                        replacement_value = f"[Invalid Date Format: {key}]"
                else:
                    replacement_value = work_data_provider.get_data(key.upper())
                    
                    if key.upper() == "TENDER_COST":
                        replacement_value = format_currency_inr(replacement_value)

                    if replacement_value is None or "[Invalid" in str(replacement_value):
                        replacement_value = evaluate_special_placeholder(key, data)

            # Handle <<PLACEHOLDER>>
            elif match.group(3):
                if is_firm_specific:
                    current_firm_name = data.get('firm_name')
                    if current_firm_name:
                        firm_document_data = work_data_provider.get_firm_document_data(current_firm_name)
                        lookup_key = key.lower()

                        if lookup_key == 'firm_name':
                            replacement_value = current_firm_name
                        elif lookup_key == 'firm_address':
                            # Get firm address from firms table
                            for firm_data in work_data_provider.firms_data:
                                if firm_data.get('name') == current_firm_name:
                                    replacement_value = firm_data.get('address')
                                    break
                        elif lookup_key == 'firm_representative':
                            # Get firm representative from firms table
                            for firm_data in work_data_provider.firms_data:
                                if firm_data.get('name') == current_firm_name:
                                    replacement_value = firm_data.get('representative')
                                    break
                        elif lookup_key == 'pg_submitted':
                            replacement_value = "submitted the PG No." if firm_document_data.get('pg_submitted') == 1 else "did not submit the PG"
                        elif lookup_key == 'indemnity_bond_submitted':
                            replacement_value = "submitted the Indemnity Bond" if firm_document_data.get('indemnity_bond_submitted') == 1 else "did not submit the Indemnity Bond"
                        elif firm_document_data and lookup_key in firm_document_data:
                            value = firm_document_data[lookup_key]
                            if lookup_key == 'pg_amount':
                                replacement_value = format_currency_inr(value)
                            else:
                                replacement_value = value
                else:
                    replacement_value = work_data_provider.get_data(key.upper())

            if replacement_value is not None:
                replacements[placeholder_full] = str(replacement_value)

        # Apply all replacements to the paragraph
        for placeholder, value in replacements.items():
            # We need to iterate through the runs and replace the text
            # This is complex because a placeholder can span multiple runs
            # A simpler approach for now is to clear the paragraph and add a new run with the replaced text
            if placeholder in paragraph.text:
                # This is a simplified replacement and might not preserve formatting perfectly across complex placeholders.
                # For more robust replacement, a run-level replacement logic is needed.
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    def _process_table_insertions(self, document):
        """Process all table insertions after text replacements are complete"""
        for insertion in self._table_insertions:
            paragraph = insertion['paragraph']
            work_id = insertion['work_id']
            reference_firm = insertion['reference_firm']
            
            # Check if the paragraph contains the table marker
            if "<<TABLE_INSERT_MARKER>>" in paragraph.text:
                # Clear the marker from the paragraph
                paragraph.text = paragraph.text.replace("<<TABLE_INSERT_MARKER>>", "")
                
                # Insert the table after this paragraph
                self.enquiry_table_formatter.create_enquiry_table(
                    document, work_id, reference_firm
                )
