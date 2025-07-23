from docx import Document
import re
from features.AutodocGen.constants import USER_PLACEHOLDER_PATTERN, WORK_DATA_PLACEHOLDER_PATTERN, FIRM_PLACEHOLDER_PATTERN, ALL_FIRMS_PG_DETAILS_PATTERN

class PlaceholderParser:
    def __init__(self):
        pass

    def extract_placeholders(self, doc_path):
        document = Document(doc_path)
        user_input_placeholders = set()
        work_data_placeholders = set()
        firm_placeholders = set()
        all_firms_pg_details_placeholders = set()

        def _extract_from_text(text_content):
            user_input_placeholders.update(re.findall(USER_PLACEHOLDER_PATTERN, text_content))
            work_data_placeholders.update(re.findall(WORK_DATA_PLACEHOLDER_PATTERN, text_content))
            firm_placeholders.update(re.findall(FIRM_PLACEHOLDER_PATTERN, text_content))
            all_firms_pg_details_placeholders.update(re.findall(ALL_FIRMS_PG_DETAILS_PATTERN, text_content))

        for paragraph in document.paragraphs:
            _extract_from_text(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _extract_from_text(paragraph.text)

        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                _extract_from_text(paragraph.text)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(paragraph.text)

            footer = section.footer
            for paragraph in footer.paragraphs:
                _extract_from_text(paragraph.text)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(paragraph.text)

        return user_input_placeholders, work_data_placeholders, firm_placeholders, all_firms_pg_details_placeholders
