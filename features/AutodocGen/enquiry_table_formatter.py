from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database.db_manager import get_schedule_items, get_firm_rate_for_item
from utils.helpers import format_currency_inr

class EnquiryTableFormatter:
    def __init__(self):
        pass
    
    def create_enquiry_table(self, document, work_id, reference_firm_name=None):
        """
        Creates an enquiry table in the document with the specified structure:
        - SN, Schedule Items, Qty, ELS KYN Estimate (merged), Firm Quoted (merged)
        - Under ELS KYN Estimate: Unit Rate, Total Cost
        - Under Firm Quoted: Unit Rate, Total Cost
        - Summary rows: GST @18% and Grand Total
        """
        
        # Get schedule items for the work
        schedule_items = get_schedule_items(work_id)
        
        if not schedule_items:
            # If no schedule items, add a simple message
            p = document.add_paragraph("No schedule items found for this work.")
            return
        
        # Create table with 8 columns for proper structure (added Unit column)
        # Header rows (2) + data rows + summary rows (2)
        num_rows = 2 + len(schedule_items) + 2
        table = document.add_table(rows=num_rows, cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(0.5)  # SN
        table.columns[1].width = Inches(2.5)  # Schedule Items  
        table.columns[2].width = Inches(0.7)  # Qty
        table.columns[3].width = Inches(0.6)  # Unit
        table.columns[4].width = Inches(1.0)  # ELS Unit Rate
        table.columns[5].width = Inches(1.2)  # ELS Total Cost
        table.columns[6].width = Inches(1.0)  # Firm Unit Rate
        table.columns[7].width = Inches(1.2)  # Firm Total Cost
        
        # Header Row 1 - Main headers
        header_row1 = table.rows[0]
        header_row1.cells[0].text = "SN"
        header_row1.cells[1].text = "Schedule Items"
        header_row1.cells[2].text = "Qty"
        header_row1.cells[3].text = "Unit"
        header_row1.cells[4].text = "ELS KYN Estimate"
        header_row1.cells[6].text = "Firm Quoted"
        
        # Merge cells for main headers
        header_row1.cells[4].merge(header_row1.cells[5])
        header_row1.cells[6].merge(header_row1.cells[7])
        
        # Header Row 2 - Sub headers
        header_row2 = table.rows[1]
        header_row2.cells[0].text = ""  # SN continues
        header_row2.cells[1].text = ""  # Schedule Items continues
        header_row2.cells[2].text = ""  # Qty continues
        header_row2.cells[3].text = ""  # Unit continues
        header_row2.cells[4].text = "Unit Rate"
        header_row2.cells[5].text = "Total Cost"
        header_row2.cells[6].text = "Unit Rate"
        header_row2.cells[7].text = "Total Cost"
        
        # Center align all header cells
        for row_idx in [0, 1]:
            for cell in table.rows[row_idx].cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Make text bold
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        
        # Data rows
        total_els_cost = 0
        
        for idx, item in enumerate(schedule_items):
            row_idx = idx + 2  # Start after header rows
            data_row = table.rows[row_idx]
            
            # SN
            data_row.cells[0].text = str(idx + 1)
            data_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Schedule Items
            data_row.cells[1].text = item.get('item_name', '')
            
            # Quantity
            qty = item.get('quantity', 0)
            data_row.cells[2].text = str(qty)
            data_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Unit
            unit = item.get('unit', '')
            data_row.cells[3].text = unit
            data_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # ELS KYN Estimate (get rate for reference firm)
            unit_rate = 0
            if reference_firm_name:
                firm_rate_data = get_firm_rate_for_item(item['item_id'], reference_firm_name)
                if firm_rate_data:
                    unit_rate = float(firm_rate_data.get('rate', 0))
            
            # ELS Unit Rate
            data_row.cells[4].text = format_currency_inr(unit_rate) if unit_rate > 0 else "₹ 0.00/-"
            data_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # ELS Total Cost
            total_cost = unit_rate * qty
            total_els_cost += total_cost
            data_row.cells[5].text = format_currency_inr(total_cost) if total_cost > 0 else "₹ 0.00/-"
            data_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Firm Quoted columns (empty for now)
            data_row.cells[6].text = ""
            data_row.cells[7].text = ""
        
        # Summary rows
        summary_start_row = len(schedule_items) + 2
        
        # GST @18% row
        gst_row = table.rows[summary_start_row]
        gst_row.cells[0].text = ""
        gst_row.cells[1].text = ""
        gst_row.cells[2].text = ""
        gst_row.cells[3].text = ""
        gst_row.cells[4].text = "GST @18%"
        gst_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        gst_amount = total_els_cost * 0.18
        gst_row.cells[5].text = format_currency_inr(gst_amount)
        gst_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        gst_row.cells[6].text = ""
        gst_row.cells[7].text = ""
        
        # Grand Total row
        total_row = table.rows[summary_start_row + 1]
        total_row.cells[0].text = ""
        total_row.cells[1].text = ""
        total_row.cells[2].text = ""
        total_row.cells[3].text = ""
        total_row.cells[4].text = "Grand Total"
        total_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        grand_total = total_els_cost + gst_amount
        total_row.cells[5].text = format_currency_inr(grand_total)
        total_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        total_row.cells[6].text = ""
        total_row.cells[7].text = ""
        
        # Make summary row text bold
        for cell in [gst_row.cells[4], gst_row.cells[5], total_row.cells[4], total_row.cells[5]]:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        
        return table
