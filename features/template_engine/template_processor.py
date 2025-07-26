from docx import Document
import re
from features.template_engine.special_placeholder_handler import evaluate_special_placeholder
from .work_data_provider import WorkDataProvider

class TemplateProcessor:
    def __init__(self):
        pass

    def extract_placeholders(self, doc_path):
        document = Document(doc_path)
        user_input_placeholders = set()

        def _extract_from_text(text_content):
            # Find all user input placeholders in the text content
            found_user_input = re.findall(r"\{\{([a-zA-Z0-9_.]+)\}\}", text_content)
            for p in found_user_input:
                user_input_placeholders.add(p)
            
        # Extract from paragraphs in the main body
        for paragraph in document.paragraphs:
            _extract_from_text(paragraph.text)

        # Extract from tables in the main body
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _extract_from_text(cell.text)

        # Extract from headers and footers
        for section in document.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                _extract_from_text(paragraph.text)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(cell.text)

            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                _extract_from_text(paragraph.text)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _extract_from_text(cell.text)

        # Filter for base placeholders only
        base_placeholders_for_gui = set()
        for p_name in user_input_placeholders:
            is_derived = False
            if p_name.startswith(("COST", "COSTAMC", "COSTRP", "COSTCON", "COSTCAMC")):
                if re.search(r"_([0-9.]+)$|_IN_WORDS$|_0$|_00$", p_name):
                    is_derived = True
            
            if p_name.startswith("DATE"):
                is_derived = False

            if not is_derived:
                base_placeholders_for_gui.add(p_name)
        
        return base_placeholders_for_gui

    def generate_letters_for_firms(self, doc_path, data, work_id, output_path):
        work_data_provider = WorkDataProvider(work_id)
        firm_names = work_data_provider.get_firm_names_list().split(', ')
        if not firm_names:
            return False, "No firms found for this work."

        master_document = Document()
        for i, firm_name in enumerate(firm_names):
            # Create a copy of the template for each firm
            firm_document = Document(doc_path)
            firm_data = data.copy()
            firm_data['firm_name'] = firm_name # Add firm name to data

            # Process all placeholders for this firm
            for paragraph in firm_document.paragraphs:
                self._replace_in_paragraph(paragraph, firm_data, work_data_provider)
            for table in firm_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, firm_data, work_data_provider)
            # ... (headers and footers)

            # Add the processed letter to the master document
            for element in firm_document.element.body:
                master_document.element.body.append(element)
            
            # Add a page break after each letter except the last one
            if i < len(firm_names) - 1:
                master_document.add_page_break()

        master_document.save(output_path)
        return True, "Letters generated successfully."

    def replace_placeholders(self, doc_path, data, work_id, output_path, firm_placeholders):
        if firm_placeholders:
            return self.generate_letters_for_firms(doc_path, data, work_id, output_path)

        document = Document(doc_path)
        work_data_provider = WorkDataProvider(work_id)

        # Process paragraphs in the main body
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, data, work_data_provider)

        # Process tables in the main body
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, work_data_provider)

        # Process headers and footers
        for section in document.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, data, work_data_provider)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data, work_data_provider)

            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_in_paragraph(paragraph, data, work_data_provider)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data, work_data_provider)

        document.save(output_path)
        return True, "Document generated successfully."

    def _replace_in_paragraph(self, paragraph, data, work_data_provider, firm_name=None):
        full_text = "".join([run.text for run in paragraph.runs])
        
        # Regex for user input placeholders only
        matches = list(re.finditer(r"\{\{([a-zA-Z0-9_.]+)\}\}", full_text))

        if not matches:
            return

        new_runs_data = []
        last_idx = 0

        for match in matches:
            placeholder_full = match.group(0)
            user_input_ph = match.group(1) # Correctly access the captured group

            if match.start() > last_idx:
                new_runs_data.append({'text': full_text[last_idx:match.start()], 'style': None})

            replacement_value = evaluate_special_placeholder(user_input_ph, data)
            if replacement_value is None or str(replacement_value).strip() == "":
                new_runs_data.append({'text': placeholder_full, 'style': None})
            else:
                new_runs_data.append({'text': str(replacement_value), 'style': None})

            last_idx = match.end()

        if last_idx < len(full_text):
            new_runs_data.append({'text': full_text[last_idx:], 'style': None})

        # Clear existing runs and add new ones
        for i in reversed(range(len(paragraph.runs))):
            p = paragraph.runs[i]
            p.element.getparent().remove(p.element)

        for run_data in new_runs_data:
            new_run = paragraph.add_run(run_data['text'])
            if run_data['style']:
                new_run.style = run_data['style']
