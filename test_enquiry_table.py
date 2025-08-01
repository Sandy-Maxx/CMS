#!/usr/bin/env python3

import sys
import os
from docx import Document

# Add current directory to Python path
sys.path.append('.')

from features.AutodocGen.data_fetcher import DataFetcher
from features.AutodocGen.document_generator import DocumentGenerator
from config import DATABASE_PATH

def test_enquiry_table():
    """Test that ENQUIRY_TABLE placeholder creates a proper table"""
    
    print("Testing ENQUIRY_TABLE placeholder...")
    
    # Create a test template with the ENQUIRY_TABLE placeholder
    test_template = Document()
    test_template.add_paragraph("Enquiry Details:")
    test_template.add_paragraph("[ENQUIRY_TABLE]")
    test_template.add_paragraph("End of enquiry table.")
    test_template.save("test_enquiry_table_template.docx")
    
    try:
        data_fetcher = DataFetcher(DATABASE_PATH)
        generator = DocumentGenerator(data_fetcher)
        
        # Test data
        test_data = {
            'work_id': 1,  # Use existing work ID
        }
        
        # Template and output paths
        template_path = 'test_enquiry_table_template.docx'
        output_path = 'test_enquiry_table_output.docx'
        
        print(f"Using template: {template_path}")
        print(f"Output will be saved to: {output_path}")
        print(f"Test data: {test_data}")
        
        # Generate the document
        generator.generate(template_path, test_data, output_path, is_firm_specific=False)
        
        print("Document generation completed!")
        
        # Read the output to verify the table was created
        print("\n--- OUTPUT DOCUMENT CONTENT ---")
        try:
            output_doc = Document(output_path)
            
            print(f"Number of paragraphs: {len(output_doc.paragraphs)}")
            print(f"Number of tables: {len(output_doc.tables)}")
            
            for i, paragraph in enumerate(output_doc.paragraphs):
                text = paragraph.text
                if text.strip():
                    print(f"Paragraph {i}: {repr(text)}")
            
            # Check if tables were created
            for i, table in enumerate(output_doc.tables):
                print(f"\nTable {i} structure:")
                print(f"  Rows: {len(table.rows)}")
                print(f"  Columns: {len(table.columns)}")
                
                # Show first few rows
                for row_idx, row in enumerate(table.rows[:5]):  # Show first 5 rows
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    print(f"  Row {row_idx}: {row_text}")
                
                if len(table.rows) > 5:
                    print(f"  ... and {len(table.rows) - 5} more rows")
                    
        except Exception as e:
            print(f"Error reading output document: {e}")
        
    except Exception as e:
        print(f"Error during test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_enquiry_table()
