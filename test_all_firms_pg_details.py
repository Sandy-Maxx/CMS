#!/usr/bin/env python3

import sys
import os
from docx import Document

# Add current directory to Python path
sys.path.append('.')

from features.AutodocGen.data_fetcher import DataFetcher
from features.AutodocGen.document_generator import DocumentGenerator
from config import DATABASE_PATH

def test_all_firms_pg_details():
    """Test that ALL_FIRMS_PG_DETAILS includes all firms, including those without documents"""
    
    print("Testing ALL_FIRMS_PG_DETAILS placeholder...")
    
    # Create a test template with the ALL_FIRMS_PG_DETAILS placeholder
    test_template = Document()
    test_template.add_paragraph("PG Details for all firms:")
    test_template.add_paragraph("[ALL_FIRMS_PG_DETAILS]")
    test_template.save("test_all_firms_template.docx")
    
    try:
        data_fetcher = DataFetcher(DATABASE_PATH)
        generator = DocumentGenerator(data_fetcher)
        
        # Test data
        test_data = {
            'work_id': 1,  # Use existing work ID
            'firm_name': 'Test Firm'
        }
        
        # Template and output paths
        template_path = 'test_all_firms_template.docx'
        output_path = 'test_all_firms_output.docx'
        
        print(f"Using template: {template_path}")
        print(f"Output will be saved to: {output_path}")
        print(f"Test data: {test_data}")
        
        # Generate the document
        generator.generate(template_path, test_data, output_path, is_firm_specific=False)
        
        print("Document generation completed!")
        
        # Read the output to see the results
        print("\n--- OUTPUT DOCUMENT CONTENT ---")
        try:
            output_doc = Document(output_path)
            for i, paragraph in enumerate(output_doc.paragraphs):
                text = paragraph.text
                if text.strip():
                    print(f"Paragraph {i}: {repr(text)}")
                    
        except Exception as e:
            print(f"Error reading output document: {e}")
        
    except Exception as e:
        print(f"Error during test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_all_firms_pg_details()
