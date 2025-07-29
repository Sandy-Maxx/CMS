import os
from database import db_manager
from features.template_engine.template_processor import TemplateProcessor
from features.template_engine.data_manager import TemplateDataManager
from docx import Document

def test_template_engine():
    print("--- Testing Template Engine ---")

    # 1. Add a test work
    print("Step 1: Adding a test work...")
    work_id = db_manager.add_work(name="Template Engine Test Work", description="Test work for template engine.")
    if not work_id:
        print("FAILURE: Could not create work")
        return
    print(f"SUCCESS: Work added with ID: {work_id}")

    # 2. Setup template and data
    print("\nStep 2: Setting up template and data...")
    template_path = "test_template.docx"
    data_manager = TemplateDataManager()
    template_processor = TemplateProcessor()
    
    # Create a minimal template document.
    try:
        doc = Document()
        doc.add_paragraph("This is a test template.")
        doc.save(template_path)
        print(f"SUCCESS: Template created at {template_path}")
    except Exception as e:
        print(f"ERROR: Unable to create template: {e}")
        return
    
    # 3. Save some initial data
    print("\nStep 3: Saving initial data...")
    initial_data = {"NAME": "Test User"}
    try:
        data_manager.save_template_data(template_path, initial_data)
        print("SUCCESS: Initial data saved.")
    except Exception as e:
        print(f"ERROR: Failed to save initial data: {e}")

    # 4. Load the data
    print("\nStep 4: Loading data...")
    loaded_data = data_manager.load_template_data(template_path)
    # Extract the 'current' value for the 'NAME' placeholder
    name_value = loaded_data.get("NAME", {}).get("current")
    if name_value == "Test User":
        print(f"SUCCESS: Loaded data: {loaded_data}")
    else:
        print(f"FAILURE: Loaded data is incorrect: {loaded_data}")

    # 5. Process the template
    print("\nStep 5: Processing the template...")
    output_path = "test_output.docx"
    # The `replace_placeholders` method expects a dictionary of placeholder values.
    # We'll pass the `initial_data` directly.
    try:
        success, message = template_processor.replace_placeholders(template_path, initial_data, work_id, output_path, firm_placeholders=set())

        if success:
            print(f"SUCCESS: {message}")
            if os.path.exists(output_path):
                print(f"SUCCESS: Output file created at: {output_path}")
                # TODO: Add verification of the content of the output file
                os.remove(output_path)
            else:
                print("FAILURE: Output file not created.")
        else:
            print(f"FAILURE: {message}")
    except Exception as e:
        print(f"ERROR: Failed to process template: {e}")

    # 6. Clean up the database and dummy files
    print("\nStep 6: Cleaning up...")
    db_manager.delete_work(work_id)
    os.remove(template_path)
    # Clean up the generated json file
    data_file = data_manager._get_template_data_path(template_path)
    if os.path.exists(data_file):
        os.remove(data_file)
    print("SUCCESS: Test work and dummy files deleted.")

def test_special_placeholder_merge():
    """Test merge output for special placeholders."""
    print("\n--- Testing Special Placeholder Merge Output ---")
    
    import tempfile
    import sqlite3
    from features.template_engine.template_processor import TemplateProcessor
    from features.template_engine.work_data_provider import WorkDataProvider
    from config import DATABASE_PATH
    import config
    import database.managers.database_utils as db_utils
    
    # Create a temporary database with test data
    with tempfile.NamedTemporaryFile(suffix='.db', delete=False) as temp_db:
        temp_db_path = temp_db.name
    
    try:
        # Set up test database
        conn = sqlite3.connect(temp_db_path)
        cursor = conn.cursor()
        
        # Create works table with test_col
        cursor.execute("""
            CREATE TABLE works (
                id INTEGER PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT,
                test_col TEXT
            )
        """)
        
        # Create firm_documents table
        cursor.execute("""
            CREATE TABLE firm_documents (
                id INTEGER PRIMARY KEY,
                work_id INTEGER,
                firm_name TEXT,
                pg_no TEXT,
                pg_amount REAL,
                bank_name TEXT,
                bank_address TEXT,
                submission_date TEXT,
                pg_submitted INTEGER
            )
        """)
        
        # Insert test data
        cursor.execute(
            "INSERT INTO works (name, description, test_col) VALUES (?, ?, ?)",
            ("Test Work", "Test Description", "Test Column Value")
        )
        
        cursor.execute(
            "INSERT INTO firm_documents (work_id, firm_name, pg_no, pg_amount, bank_name, bank_address, submission_date, pg_submitted) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (1, "Test Firm", "PG123", 50000.0, "Test Bank", "Test Address", "2024-01-01", 1)
        )
        
        conn.commit()
        conn.close()
        
        # Create a test template document
        test_doc = Document()
        paragraph = test_doc.add_paragraph()
        # Add various placeholder types
        paragraph.add_run("Work Name: [NAME]\n")
        paragraph.add_run("Work Description: [DESCRIPTION]\n")
        paragraph.add_run("Test Column: [TEST_COL]\n")
        paragraph.add_run("Current Date: [CURRENT_DATE]\n")
        paragraph.add_run("Current Time: [CURRENT_TIME]\n")
        paragraph.add_run("Firm Name: <<FIRM_NAME>>\n")
        paragraph.add_run("PG Number: <<PG_NO>>\n")
        paragraph.add_run("User Input: {{TEST_INPUT}}\n")
        
        template_path = "test_merge_template.docx"
        test_doc.save(template_path)
        
        # Redirect database to test database
        original_database_path = DATABASE_PATH
        config.DATABASE_PATH = temp_db_path
        db_utils.DATABASE_PATH = temp_db_path
        
        try:
            # Test template processor
            processor = TemplateProcessor()
            output_path = "test_merge_output.docx"
            
            # Test user input data
            user_data = {"TEST_INPUT": "User Input Value"}
            
            # Process the template
            success, message = processor.replace_placeholders(
                template_path, user_data, 1, output_path, firm_placeholders=set()
            )
            
            if success:
                print(f"SUCCESS: Template processing completed - {message}")
                
                # Verify output document content
                output_doc = Document(output_path)
                output_text = ""
                for paragraph in output_doc.paragraphs:
                    output_text += paragraph.text + "\n"
                
                print(f"Output document content:\n{output_text}")
                
                # Check for specific replacements
                checks = [
                    ("Test Work" in output_text, "Work name replacement"),
                    ("Test Description" in output_text, "Work description replacement"),
                    ("Test Column Value" in output_text, "TEST_COL placeholder replacement"),
                    ("Test Firm" in output_text, "Firm name replacement"),
                    ("PG123" in output_text, "PG number replacement"),
                    ("User Input Value" in output_text, "User input replacement"),
                    ("[CURRENT_DATE]" not in output_text, "Current date replacement"),
                    ("[CURRENT_TIME]" not in output_text, "Current time replacement")
                ]
                
                for check, description in checks:
                    if check:
                        print(f"SUCCESS: {description}")
                    else:
                        print(f"FAILURE: {description}")
                
                # Clean up output file
                if os.path.exists(output_path):
                    os.remove(output_path)
            else:
                print(f"FAILURE: Template processing failed - {message}")
                
        finally:
            # Restore original database path
            config.DATABASE_PATH = original_database_path
            db_utils.DATABASE_PATH = original_database_path
            
        # Clean up template file
        if os.path.exists(template_path):
            os.remove(template_path)
            
    finally:
        # Clean up temporary database
        if os.path.exists(temp_db_path):
            os.unlink(temp_db_path)

def test_gui_about_tab_placeholder_display():
    """Test that About tab shows new placeholder in the list."""
    print("\n--- Testing GUI About Tab Placeholder Display ---")
    
    import tempfile
    import sqlite3
    from features.about_tab.about_tab import AboutTab
    from config import DATABASE_PATH
    import config
    import database.managers.database_utils as db_utils
    import tkinter as tk
    
    # Create a temporary database with test_col
    with tempfile.NamedTemporaryFile(suffix='.db', delete=False) as temp_db:
        temp_db_path = temp_db.name
    
    try:
        # Set up test database
        conn = sqlite3.connect(temp_db_path)
        cursor = conn.cursor()
        
        # Create works table with test_col
        cursor.execute("""
            CREATE TABLE works (
                id INTEGER PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT,
                test_col TEXT
            )
        """)
        
        # Create firm_documents table
        cursor.execute("""
            CREATE TABLE firm_documents (
                id INTEGER PRIMARY KEY,
                work_id INTEGER,
                firm_name TEXT
            )
        """)
        
        conn.commit()
        conn.close()
        
        # Redirect database to test database
        original_database_path = DATABASE_PATH
        config.DATABASE_PATH = temp_db_path
        db_utils.DATABASE_PATH = temp_db_path
        
        try:
            # Create a minimal Tkinter window for testing
            root = tk.Tk()
            root.withdraw()  # Hide the window
            
            # Create AboutTab instance
            about_tab = AboutTab(root)
            
            # Get the dynamic placeholder content
            placeholder_content = about_tab._get_dynamic_placeholder_content()
            
            if "[TEST_COL]" in placeholder_content:
                print("SUCCESS: [TEST_COL] appears in About tab placeholder list")
            else:
                print("FAILURE: [TEST_COL] not found in About tab placeholder list")
                print(f"Placeholder content preview: {placeholder_content[:500]}...")
            
            root.destroy()
            
        finally:
            # Restore original database path
            config.DATABASE_PATH = original_database_path
            db_utils.DATABASE_PATH = original_database_path
            
    finally:
        # Clean up temporary database
        if os.path.exists(temp_db_path):
            os.unlink(temp_db_path)

if __name__ == "__main__":
    db_manager.create_tables()
    test_template_engine()
    test_special_placeholder_merge()
    test_gui_about_tab_placeholder_display()
